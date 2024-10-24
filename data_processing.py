'''
import os
import json
import PyPDF2
import docx
import pytesseract
from PIL import Image
import pandas as pd

# Process PDF Files
def process_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    return {"content": text}

# Process Word Files (.docx)
def process_word(file_path):
    doc = docx.Document(file_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return {"content": text}

# Process Excel Files
def process_excel(file_path):
    df = pd.read_excel(file_path)
    return df.to_dict(orient="records")

# Process Image Files (JPEG, JPG, PNG) using OCR
def process_image(file_path):
    image = Image.open(file_path)
    text = pytesseract.image_to_string(image)
    return {"content": text}

# Main function to determine file type and process
def process_file(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    if extension == '.pdf':
        return process_pdf(file_path)
    elif extension == '.docx':
        return process_word(file_path)
    elif extension in ['.xlsx', '.xls']:
        return process_excel(file_path)
    elif extension in ['.jpg', '.jpeg', '.png']:
        return process_image(file_path)
    else:
        return {"error": "Unsupported file format"}
'''
import os
import json
import pandas as pd
import PyPDF2
import docx
import pytesseract
from PIL import Image
import zipfile

# Process PDF Files
def process_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
    return {"content": text}

# Process Word Files (.docx)
def process_word(file_path):
    doc = docx.Document(file_path)
    text = '\n'.join([para.text for para in doc.paragraphs])
    return {"content": text}

# Process Excel Files (.xls, .xlsx)
def process_excel(file_path):
    df = pd.read_excel(file_path)
    return df.to_dict(orient="records")

# Process CSV Files (.csv)
def process_csv(file_path):
    df = pd.read_csv(file_path)
    return df.to_dict(orient="records")

# Process Image Files (JPEG, JPG, PNG) using OCR
def process_image(file_path):
    image = Image.open(file_path)
    text = pytesseract.image_to_string(image)
    return {"content": text}

# Process ZIP Files
def process_zip(file_path):
    extracted_data = []
    with zipfile.ZipFile(file_path, 'r') as zip_ref:
        zip_ref.extractall('uploads/')  # Extract to uploads directory
        for extracted_file in zip_ref.namelist():
            extracted_file_path = os.path.join('uploads/', extracted_file)
            if extracted_file.endswith('.pdf'):
                extracted_data.append(process_pdf(extracted_file_path))
            elif extracted_file.endswith('.docx'):
                extracted_data.append(process_word(extracted_file_path))
            elif extracted_file.endswith('.xlsx') or extracted_file.endswith('.xls'):
                extracted_data.append(process_excel(extracted_file_path))
            elif extracted_file.endswith('.csv'):
                extracted_data.append(process_csv(extracted_file_path))
            elif extracted_file.endswith(('.jpg', '.jpeg', '.png')):
                extracted_data.append(process_image(extracted_file_path))
    return {"files": extracted_data}

# Main function to determine file type and process
def process_file(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    if extension == '.pdf':
        return process_pdf(file_path)
    elif extension == '.docx':
        return process_word(file_path)
    elif extension in ['.xlsx', '.xls']:
        return process_excel(file_path)
    elif extension == '.csv':
        return process_csv(file_path)
    elif extension in ['.jpg', '.jpeg', '.png']:
        return process_image(file_path)
    elif extension == '.zip':
        return process_zip(file_path)
    else:
        return {"error": "Unsupported file format"}
